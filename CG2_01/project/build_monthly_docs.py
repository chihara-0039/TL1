from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "monthly_submission_docs"
OUT.mkdir(exist_ok=True)

ACCOUNT = "LE3C_15_チハラ_シゴウ"
WORK_TITLE = "トラブル・シューティング"
PROJECT_PATH = r"01_作品_トラブル・シューティング\DirectXGame"
RELEASE_PATH = r"01_作品_トラブル・シューティング\リリースファイル"
VIDEO_PATH = r"04_作品紹介動画\トラブルシューティング_デモ動画.mp4"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(90, 97, 110)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"


def set_font(run, size=11, bold=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_title(doc, title, subtitle, size=24):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=size, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_font(run, size=12, color=MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(14)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D7DBE2")
    border.append(bottom)
    rule._p.get_or_add_pPr().append(border)


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_font(run, size=11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, size=11)


def add_kv_table(doc, rows, widths=(1.55, 4.95)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), label, bold=True, fill=BLUE_FILL)
        set_cell_text(table.cell(idx, 1), value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, fill=GRAY_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for s in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[s].font.name = "Calibri"
        styles[s]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    return doc


def build_portfolio():
    doc = setup_doc()
    add_title(doc, f"{WORK_TITLE} ポートフォリオ", f"{ACCOUNT} / 3Dボスシューティング作品")
    add_kv_table(doc, [
        ("作品名", WORK_TITLE),
        ("ジャンル", "3Dシューティング / ボス戦アクション"),
        ("制作環境", "C++ / DirectX / KamataEngine / Visual Studio 2022"),
        ("提出物", "ビルド済みRelease、作品紹介動画、ビルド可能なソースコード一式、プログラム説明資料"),
    ])

    add_heading(doc, "作品概要")
    add_body(doc, "プレイヤー機体を操作し、画面奥に出現するボスを撃破する3Dシューティングゲームです。通常射撃に加え、照準方向へ弾を撃つ操作、ボスのファンネル攻撃、HP表示、クリア/ゲームオーバー遷移までを一連のゲーム体験として実装しました。")

    add_heading(doc, "見てほしいポイント")
    add_bullets(doc, [
        "ボスの周囲を回る3基のファンネルが、プレイヤーを狙って連続射撃する攻撃パターン",
        "プレイヤーとボスのHPを画面上に表示し、戦況が分かりやすいHUDにした点",
        "タイトル、ゲーム本編、クリア、ゲームオーバーをシーン管理で切り替える構成",
        "OBJモデル、テクスチャ、効果音、BGMを組み合わせ、宇宙空間のボス戦らしい雰囲気を作った点",
    ])

    add_heading(doc, "操作とゲームルール")
    add_matrix(doc, ["項目", "内容"], [
        ("移動", "W/A/S/Dで上下左右に移動"),
        ("攻撃", "照準方向へ射撃し、ボスにダメージを与える"),
        ("勝利条件", "ボスのHPを0にするとクリア画面へ遷移"),
        ("敗北条件", "プレイヤーHPが0になるとゲームオーバー画面へ遷移"),
    ], [1.45, 5.05])

    add_heading(doc, "担当・実装範囲")
    add_bullets(doc, [
        "プレイヤー移動、射撃、被弾時の無敵時間、HP管理",
        "ボス本体の行動、HP管理、被弾フラッシュ、撃破判定",
        "ファンネル本体の追従移動と弾幕生成",
        "弾の生成、更新、寿命管理、当たり判定",
        "シーン管理、HUD、BGM/効果音、モデル/テクスチャ読み込み",
    ])

    add_heading(doc, "今後伸ばしたい点")
    add_bullets(doc, [
        "攻撃パターンを増やし、ボスのHP段階に応じて難度が変わる構成にする",
        "プレイヤー側のスキル、回避、チャージ攻撃などを追加して駆け引きを増やす",
        "エフェクトとカメラ演出を強化し、PVで伝わる派手さを上げる",
    ])

    path = OUT / f"{ACCOUNT}_ポートフォリオ.docx"
    doc.save(path)
    return path


def build_program_doc():
    doc = setup_doc()
    add_title(doc, f"{WORK_TITLE} プログラム説明資料", f"{ACCOUNT} / C++ DirectX 3Dシューティング", size=22)
    add_kv_table(doc, [
        ("対象作品", WORK_TITLE),
        ("プロジェクト", PROJECT_PATH),
        ("Release", RELEASE_PATH),
        ("動画", VIDEO_PATH),
        ("主な技術", "C++20、DirectX、KamataEngine、OBJモデル描画、シーン管理、当たり判定、サウンド再生"),
    ])

    add_heading(doc, "プログラム構成")
    add_matrix(doc, ["クラス/ファイル", "役割"], [
        ("SceneManager", "タイトル、ゲーム、クリア、ゲームオーバーのシーン遷移を管理"),
        ("GameScene", "ゲーム本編の初期化、更新、描画、当たり判定、勝敗判定を統括"),
        ("Player", "移動、射撃、HP、被弾時の無敵時間、プレイヤー弾の管理"),
        ("Enemy", "ボスの移動、HP、被弾処理、攻撃パターンの基礎処理"),
        ("Funnel / FunnelBullet", "ボス周囲を回る支援機と、プレイヤーを狙う弾を管理"),
        ("HpHud", "プレイヤーHPとボスHPを画面上に表示"),
        ("Skydome", "背景空間を描画し、3Dシューティングの舞台を作る"),
    ], [1.75, 4.75])

    add_heading(doc, "更新処理の流れ")
    add_bullets(doc, [
        "入力を取得し、プレイヤーの移動と射撃を更新する",
        "ボス本体とファンネルの位置を更新し、攻撃パターンに応じて弾を生成する",
        "プレイヤー弾、敵弾を更新し、寿命切れの弾を削除する",
        "プレイヤーと敵弾、ボスとプレイヤー弾の当たり判定を行う",
        "HPが0になった対象を確認し、クリアまたはゲームオーバーへシーン遷移する",
    ])

    add_heading(doc, "当たり判定とダメージ処理")
    add_body(doc, "敵弾とプレイヤーは距離ベースで判定し、プレイヤー弾とボスはAABBで判定しています。被弾時はHPを減らし、プレイヤーには短時間の無敵時間を持たせて連続ヒットを防いでいます。ボスは被弾フラッシュでダメージを視覚的に伝える構成です。")

    add_heading(doc, "ボス/ファンネル攻撃")
    add_body(doc, "ボスの周囲に3基のファンネルを配置し、ボス位置を中心に追従させます。攻撃パターンでは各ファンネルからプレイヤー方向へ弾を連射し、ボス本体だけでなく周辺ユニットも脅威になるようにしました。")

    add_heading(doc, "ビルド・実行")
    add_matrix(doc, ["項目", "内容"], [
        ("必要環境", "Windows 10/11、Visual Studio 2022、DirectX対応環境"),
        ("構成", "Release | x64"),
        ("起動", "Releaseフォルダ内のDirectXGame.exeを実行"),
        ("注意", "Resourcesフォルダをexeと同じ階層に置く。ソースをビルドする場合はDirectXGameと同じ階層のExternalも必要。"),
    ], [1.45, 5.05])

    add_heading(doc, "アピールしたい技術点")
    add_bullets(doc, [
        "ゲーム本編の流れをGameSceneにまとめ、各オブジェクトの役割をクラスに分けたこと",
        "弾やファンネルをリストで管理し、寿命切れやヒット済みを削除する更新構造",
        "HP HUD、シーン遷移、サウンド、3Dモデル描画を組み合わせ、作品として最後まで遊べる形にしたこと",
    ])

    path = OUT / f"{ACCOUNT}_プログラム説明資料.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_portfolio())
    print(build_program_doc())
